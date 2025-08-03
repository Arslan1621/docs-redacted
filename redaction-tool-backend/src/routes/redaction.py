from flask import Blueprint, request, jsonify, send_file
from werkzeug.utils import secure_filename
import os
import tempfile
import json
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import docx2txt
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph
import io

redaction_bp = Blueprint('redaction', __name__)

UPLOAD_FOLDER = tempfile.gettempdir()
ALLOWED_EXTENSIONS = {'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@redaction_bp.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        file.save(filepath)
        
        # Extract text content from the document
        try:
            doc = Document(filepath)
            content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append({
                        'type': 'paragraph',
                        'text': paragraph.text,
                        'id': len(content)
                    })
            
            return jsonify({
                'success': True,
                'filename': filename,
                'content': content
            })
        except Exception as e:
            return jsonify({'error': f'Error processing document: {str(e)}'}), 500
    
    return jsonify({'error': 'Invalid file type. Only .docx files are allowed.'}), 400

@redaction_bp.route('/redact', methods=['POST'])
def apply_redaction():
    data = request.get_json()
    filename = data.get('filename')
    redactions = data.get('redactions', [])
    
    if not filename:
        return jsonify({'error': 'No filename provided'}), 400
    
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    
    if not os.path.exists(filepath):
        return jsonify({'error': 'File not found'}), 404
    
    try:
        # Store redaction information for later use
        redaction_data = {
            'filename': filename,
            'redactions': redactions
        }
        
        redaction_file = os.path.join(UPLOAD_FOLDER, f"{filename}_redactions.json")
        with open(redaction_file, 'w') as f:
            json.dump(redaction_data, f)
        
        return jsonify({'success': True, 'message': 'Redactions applied successfully'})
    
    except Exception as e:
        return jsonify({'error': f'Error applying redactions: {str(e)}'}), 500

@redaction_bp.route('/download/<format_type>/<filename>')
def download_redacted(format_type, filename):
    if format_type not in ['docx', 'pdf']:
        return jsonify({'error': 'Invalid format. Use docx or pdf.'}), 400
    
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    redaction_file = os.path.join(UPLOAD_FOLDER, f"{filename}_redactions.json")
    
    if not os.path.exists(filepath):
        return jsonify({'error': 'Original file not found'}), 404
    
    # Load redaction data
    redactions = []
    if os.path.exists(redaction_file):
        with open(redaction_file, 'r') as f:
            redaction_data = json.load(f)
            redactions = redaction_data.get('redactions', [])
    
    try:
        if format_type == 'docx':
            return download_docx(filepath, redactions, filename)
        else:
            return download_pdf(filepath, redactions, filename)
    
    except Exception as e:
        return jsonify({'error': f'Error generating {format_type}: {str(e)}'}), 500

def download_docx(filepath, redactions, original_filename):
    doc = Document(filepath)
    
    # Apply redactions to the document
    for redaction in redactions:
        paragraph_id = redaction.get('paragraphId')
        start_pos = redaction.get('startPos', 0)
        end_pos = redaction.get('endPos', 0)
        
        if paragraph_id < len(doc.paragraphs):
            paragraph = doc.paragraphs[paragraph_id]
            original_text = paragraph.text
            
            if start_pos < len(original_text) and end_pos <= len(original_text):
                # Calculate redaction blocks needed
                redacted_length = end_pos - start_pos
                redaction_blocks = '█' * redacted_length
                
                # Replace the text
                new_text = original_text[:start_pos] + redaction_blocks + original_text[end_pos:]
                paragraph.clear()
                paragraph.add_run(new_text)
    
    # Save to temporary file
    output_filename = f"redacted_{original_filename}"
    output_path = os.path.join(UPLOAD_FOLDER, output_filename)
    doc.save(output_path)
    
    return send_file(output_path, as_attachment=True, download_name=output_filename)

def download_pdf(filepath, redactions, original_filename):
    # Extract text from original document
    text_content = docx2txt.process(filepath)
    
    # Apply redactions to text
    lines = text_content.split('\n')
    
    # Create a simple mapping for redactions (this is a simplified approach)
    for redaction in redactions:
        paragraph_id = redaction.get('paragraphId')
        start_pos = redaction.get('startPos', 0)
        end_pos = redaction.get('endPos', 0)
        
        if paragraph_id < len(lines):
            line = lines[paragraph_id]
            if start_pos < len(line) and end_pos <= len(line):
                redacted_length = end_pos - start_pos
                redaction_blocks = '█' * redacted_length
                lines[paragraph_id] = line[:start_pos] + redaction_blocks + line[end_pos:]
    
    # Create PDF
    output_filename = f"redacted_{original_filename.replace('.docx', '.pdf')}"
    output_path = os.path.join(UPLOAD_FOLDER, output_filename)
    
    doc = SimpleDocTemplate(output_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    for line in lines:
        if line.strip():
            para = Paragraph(line, styles['Normal'])
            story.append(para)
    
    doc.build(story)
    
    return send_file(output_path, as_attachment=True, download_name=output_filename)

